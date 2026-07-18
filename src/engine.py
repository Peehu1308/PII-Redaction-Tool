import re
import json
import spacy
from faker import Faker
from collections import Counter
from docx import Document
import sys
import os

# Insert the src directory at the front of sys.path so our local detectors.py
# takes precedence over any pip package also named 'detectors'.
_SRC_DIR = os.path.dirname(os.path.abspath(__file__))
if _SRC_DIR not in sys.path:
    sys.path.insert(0, _SRC_DIR)

from detectors import (
    EMAIL_RE, IP_RE, SSN_RE, DOB_RE, CIN_RE, PAN_RE, GST_RE, DIN_RE,
    find_phones, find_credit_cards, safe_replace, is_probably_real_person,
    is_probably_real_org, looks_like_address
)

fake = Faker()
Faker.seed(42)

def read_docx(path):
    doc = Document(path)
    return [p.text for p in doc.paragraphs if p.text.strip()]

class RedactionEngine:
    def __init__(self, nlp=None):
        self.nlp = nlp or spacy.load("en_core_web_sm", disable=["lemmatizer"])
        self.person_map = {}
        self.org_map = {}
        self.address_map = {}
        self.email_map = {}
        self.phone_map = {}
        self.ssn_map = {}
        self.entity_counter = Counter()

    def _fake_person(self, original):
        key = self._norm(original)
        if key not in self.person_map:
            name = fake.name()
            self.person_map[key] = self._match_case(original, name)
        return self.person_map[key]

    def _fake_org(self, original):
        key = self._norm(original)
        if key not in self.org_map:
            suffix = ""
            low = original.lower()
            for s in ["private limited", "limited", "llp", "family trust", "trust"]:
                if low.strip().endswith(s):
                    suffix = s.title()
                    break
            
            base = fake.company().split()[0].replace(",", "")
            if suffix:
                name = f"{base} {suffix}"
            else:
                name = fake.company().replace(",", "")
                
            self.org_map[key] = self._match_case(original, name)
        return self.org_map[key]

    def _fake_address(self, original):
        key = self._norm(original)
        if key not in self.address_map:
            addr = fake.address().replace("\n", ", ")
            self.address_map[key] = addr
        return self.address_map[key]

    def fake_email(self, original):
        key = self._norm(original)
        if key not in self.email_map:
            first = fake.first_name().lower()
            last = fake.last_name().lower()
            self.email_map[key] = f"{first}.{last}@example.com"
        return self.email_map[key]

    def fake_phone(self, original):
        key = self._norm(original)
        if key not in self.phone_map:
            prefix_match = re.match(r"^\s*\+\d{1,3}", original)
            prefix = prefix_match.group().strip() if prefix_match else ""
            digits_needed = max(6, len(re.sub(r"\D", "", original)) - len(re.sub(r"\D", "", prefix)))
            fake_digits = "".join(fake.random.choices("0123456789", k=digits_needed))
            self.phone_map[key] = f"{prefix} {fake_digits}".strip()
        return self.phone_map[key]

    def fake_ssn(self, original):
        key = self._norm(original)
        if key not in self.ssn_map:
            self.ssn_map[key] = fake.ssn()
        return self.ssn_map[key]

    @staticmethod
    def _norm(s):
        return re.sub(r"\s+", " ", s.strip().lower())

    @staticmethod
    def _match_case(original, replacement):
        if original.isupper():
            return replacement.upper()
        if original.istitle():
            return replacement.title()
        return replacement

    def extract_entities(self, paragraphs):
        persons, orgs, addresses = set(), set(), set()
        for doc in self.nlp.pipe(paragraphs, batch_size=64):
            for ent in doc.ents:
                text = ent.text.strip()
                if ent.label_ == "PERSON" and is_probably_real_person(text):
                    persons.add(text)
                    self.entity_counter[("PERSON", self._norm(text))] += 1
                elif ent.label_ == "ORG" and is_probably_real_org(text):
                    orgs.add(text)
                    self.entity_counter[("ORG", self._norm(text))] += 1
                elif ent.label_ in ("GPE", "LOC", "FAC"):
                    if looks_like_address(text):
                        addresses.add(text)
                        self.entity_counter[("ADDRESS", self._norm(text))] += 1
        
        expanded_persons = set(persons)
        for p in persons:
            parts = p.split()
            if len(parts) > 1 and len(parts[0]) > 2:
                expanded_persons.add(parts[0])
                
        filtered_orgs = {o for o in orgs if o not in expanded_persons}
        
        return {"PERSON": expanded_persons, "ORG": filtered_orgs, "ADDRESS": addresses}

    def build_maps(self, persons, orgs, addresses):
        for p in sorted(persons, key=len, reverse=True):
            self._fake_person(p)
        for o in sorted(orgs, key=len, reverse=True):
            self._fake_org(o)
        for a in sorted(addresses, key=len, reverse=True):
            self._fake_address(a)

    def dump_maps(self, path):
        with open(path, "w") as f:
            json.dump({
                "persons": self.person_map,
                "orgs": self.org_map,
                "addresses": self.address_map,
            }, f, indent=2, ensure_ascii=False)

    def process_document(self, input_path, output_path):
        paragraphs = read_docx(input_path)
        entities = self.extract_entities(paragraphs)
        
        self.build_maps(entities["PERSON"], entities["ORG"], entities.get("ADDRESS", set()))
        
        replacements = []
        for p in entities["PERSON"]:
            replacements.append((p, self._fake_person(p)))
        for o in entities["ORG"]:
            replacements.append((o, self._fake_org(o)))
        for a in entities.get("ADDRESS", set()):
            replacements.append((a, self._fake_address(a)))
        
        replacements.sort(key=lambda x: len(x[0]), reverse=True)
        
        def redact_text_chunk(text):
            if not text or not text.strip():
                return text
            
            text = EMAIL_RE.sub(lambda m: self.fake_email(m.group()), text)
            text = IP_RE.sub("[IP]", text)
            text = SSN_RE.sub(lambda m: self.fake_ssn(m.group()), text)
            text = DOB_RE.sub("[DOB]", text)
            text = CIN_RE.sub("[CIN]", text)
            text = PAN_RE.sub("[PAN]", text)
            text = GST_RE.sub("[GSTIN]", text)
            text = DIN_RE.sub("[DIN]", text)
            
            for _, _, phone in find_phones(text):
                text = safe_replace(text, phone, self.fake_phone(phone))
                
            for _, _, cc in find_credit_cards(text):
                text = safe_replace(text, cc, "[CREDIT CARD]")
                
            for original, fake_str in replacements:
                text = safe_replace(text, original, fake_str)
                
            return text

        doc = Document(input_path)
        
        for para in doc.paragraphs:
            original_text = para.text
            if original_text:
                new_text = redact_text_chunk(original_text)
                if new_text != original_text:
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_text
                    else:
                        para.add_run(new_text)
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        if original_text:
                            new_text = redact_text_chunk(original_text)
                            if new_text != original_text:
                                for run in para.runs:
                                    run.text = ""
                                if para.runs:
                                    para.runs[0].text = new_text
                                else:
                                    para.add_run(new_text)
                                
        doc.save(output_path)
